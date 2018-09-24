import sys
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import QThread, pyqtSignal, pyqtSlot
from PyQt5.QtWidgets import QFileDialog, QGridLayout, QProgressBar
import docx
from docx import Document
import mammoth
import bs4
from bs4 import BeautifulSoup
import json
import subprocess
import os.path
import time      
        

class Window(QtWidgets.QWidget):
    
    def __init__(self):
        super().__init__()
        self.init_ui()
        
        
    def init_ui(self):
        self.grid = QGridLayout()
        self.grid.setSpacing(10)
#        self.grid.setColumnStretch(1, 0)
#        self.grid.setColumnStretch(1, 1)
        
        
        self.butt = False
        self.newfont = QtGui.QFont("Times", 8, QtGui.QFont.Bold) 
        self.headfont = QtGui.QFont("Times", 20, QtGui.QFont.Bold) 
        self.b=QtWidgets.QPushButton('Choose File')
        self.b1=QtWidgets.QPushButton('Run')
        self.b1.setEnabled(False)
        
        self.b.setFixedSize( 100, 40 )
        self.b1.setFixedSize( 100, 40 )
        self.l = QtWidgets.QLabel('File:')
        self.h=QtWidgets.QLabel('Json Generator')
        self.l.setFont(self.newfont)
        self.h.setFont(self.headfont)
        
        self.grid.addWidget(self.h, 0, 0, 1,0)
        self.grid.addWidget(self.b, 2, 0)
        self.grid.addWidget(self.l, 2, 1)
        self.grid.addWidget(self.b1, 3, 0)
#        self.progress = QProgressBar(self)
#        self.grid.addWidget(self.progress, 4, 1)
#        self.grid.setRowStretch(0, 2)
#        self.l.move(15, 10)
        
#        h_box=QtWidgets.QHBoxLayout()
#        h_box.addStretch()
#        h_box.addWidget(self.l)
#        h_box.addStretch()
#        self.progress = QProgressBar(self)
#        self.grid.addWidget(self.progress, 3, 3)

        self.setFixedSize(500,200)
#        v_box=QtWidgets.QVBoxLayout()
#        v_box.addWidget(self.b)
#        v_box.addWidget(self.b1)
#        v_box.addLayout(h_box)
        self.thread = QThread()
        self.setLayout(self.grid)
        self.setWindowTitle('Json Generator')
        
        self.b.clicked.connect(self.selectFile)
        self.b1.clicked.connect(self.doctorPython)
        self.show()
        self.doc='hello'
        self.data = []
        
#        self.thread = QThread()
        
        
#        self.thread.started.connect(self.worker.doctorPython)
    def btn_click(self, text):
        print (text)
#        self.l.setText('clicked')
#        self.progress.setValue(text)
#        self.document.emit(self.doc)
#        self.worker.doctorPython(self.doc)
#    def using_q_thread():
#        app = QCoreApplication([])
#        thread = AThread()
#        thread.finished.connect(app.exit)
#        thread.start()
#        sys.exit(app.exec_())
    
    def selectFile(self):
        self.data = []
        self.doc=str((QFileDialog.getOpenFileName())[0])
        self.l.setText('File: '+ self.doc)
        self.b1.setEnabled(True)
        
    def doctorPython(self):
#        self.worker.moveToThread(self.thread)
        print('here')
        
#        self.worker.start()
#        self.thread.started.connect(self.worker.doctorPython)
#        self.thread.start()
#        self.worker.emit('hello there')
#        self.worker.dropped.connect(self.btn_click)
        
#        self.doctorPython(self.doc)
#        self.thread.start()
        self.b1.setEnabled(False)
        
        if(self.doc != 'hello' and self.butt == False):
            self.butt = True
            document = Document(self.doc)
            table = document.tables[0]
            with open(self.doc, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html = result.value # The generated HTML
                messages = result.messages # Any messages, such as warnings during conversion
                soup = BeautifulSoup(html, "lxml")
                print('hello hello')
            for x in range(len(document.tables)):
                print('doctor')
#                QtCore.QCoreApplication.processEvents()
                self.parseTables(document.tables[x], x, soup)
            subprocess.Popen(r'explorer /select,"data.json"') 
            open('data.json', 'w').close() 
            with open('data.json', 'w') as outfile:
                 json.dump(self.data, outfile)
            self.butt = False
            
            
    def parseTables(self, table, index, soup):

        keys = ("Field1", "Field2", "Field3", "Field4", "Field5","Field6","Field7","Field8","Field9","Field10", "Field11","Field12","Field13")
        subKeys = ("Sub-Field-1","Sub-Field-2","Sub-Field-3","Sub-Field-4","Sub-Field-5","Sub-Field-6","Sub-Field-7","Sub-Field-8","Sub-Field-9","Sub-Field-10","Sub-Field-11","Sub-Field-12","Sub-Field-13","Sub-Field-14","Sub-Field-15","Sub-Field-16","Sub-Field-17","Sub-Field-18","Sub-Field-19","Sub-Field-20","Sub-Field-21")
        for i, column in enumerate(table.columns):
            result = (mammoth.convert_to_html(cell) for cell in column.cells)
            text = (cell.text.strip() for cell in column.cells)
            if i == 0:
                continue

            if i == 2:
                continue
            row_data = dict(zip(keys, text))
            self.data.append(row_data)

        sub=[] 
#        self.progress.setValue(index)
        print(index)
        for x in range(21):
            sub.append(table.cell(13 + x,2).text)
        sub_data = dict(zip(subKeys, sub))   
#        word1 = soup.findAll("p", string="#Field9:")[index].find_next('td').contents
#        word1 = ['{0}'.format(element) for element in word1]
#        word1=''.join(word1)
#        word2 = soup.findAll("p", string="#Field10:")[index].find_next('td').contents
#        word2 = ['{0}'.format(element) for element in word2]
#        word2=''.join(word2)
#        word3 = soup.findAll("p", string="#Field11:")[index].find_next('td').contents
#        word3 = ['{0}'.format(element) for element in word3]
#        word3=''.join(word3)
#        self.data[index]['Field9'] = str(word1)
#        self.data[index]['Field10'] = str(word2)
#        self.data[index]['Field11'] = str(word3)
#        self.data[index]['Field14'] = sub_data  
        
app = QtWidgets.QApplication(sys.argv)
a_window = Window()
sys.exit(app.exec_())

